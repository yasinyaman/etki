"""Converts uploaded documents (txt/md/csv/docx/xlsx/pdf) to text + request lines.

`parse_document(filename, data) -> (full_text, items)`:
  - full_text: the full text, for scope extraction.
  - items: a list of meaningful lines/paragraphs/rows for triage (each a request candidate).
Heavy libraries (docx/openpyxl/pypdf) are lazily imported only for the relevant format.
"""

from __future__ import annotations

import csv
import io
import zipfile
from pathlib import Path

# Decompression-bomb guards. The upload byte-size cap (web._over_upload_limit) bounds only
# the COMPRESSED payload; docx/xlsx are ZIP containers, so a few-KB upload can declare/expand
# to gigabytes of XML and OOM the (single) worker. Bound the DECLARED uncompressed total and
# the expansion ratio before handing bytes to python-docx/openpyxl. PDF is bounded by page +
# extracted-text caps instead (pypdf streams, no central-directory size to pre-check).
_MAX_UNCOMPRESSED_BYTES = 512 * 1024 * 1024  # 512 MB expanded ceiling
_MAX_COMPRESSION_RATIO = 200  # reject archives expanding > 200×
_MAX_PDF_PAGES = 5_000
_MAX_PDF_TEXT_CHARS = 20 * 1024 * 1024  # 20 M chars extracted-text ceiling


class DocumentTooLarge(ValueError):
    """An upload whose decompressed size exceeds the safety bound (zip/xml bomb guard)."""


def _guard_zip_bomb(data: bytes) -> None:
    """Reject a ZIP-based document (docx/xlsx) that declares an unsafe expansion BEFORE any
    library decompresses it. Uses the central-directory `file_size` (what a real decompress
    must produce), plus a compression-ratio cap for the pathological highly-compressible case."""
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as zf:
            total = sum(info.file_size for info in zf.infolist())
    except zipfile.BadZipFile:
        return  # not a valid archive → the format parser fails loudly on its own
    if total > _MAX_UNCOMPRESSED_BYTES:
        raise DocumentTooLarge(f"açılmış boyut çok büyük ({total} bayt)")
    if data and total / len(data) > _MAX_COMPRESSION_RATIO:
        raise DocumentTooLarge(f"aşırı sıkıştırma oranı (~{total // max(len(data), 1)}×)")


def _meaningful(line: str) -> bool:
    stripped = line.strip()
    return len(stripped) > 3 and any(c.isalpha() for c in stripped)


def _from_text(text: str) -> tuple[str, list[str]]:
    items = [ln.strip() for ln in text.splitlines() if _meaningful(ln)]
    return text, items


def _from_csv(data: bytes, delimiter: str) -> tuple[str, list[str]]:
    text = data.decode("utf-8", errors="replace")
    items: list[str] = []
    for row in csv.reader(io.StringIO(text), delimiter=delimiter):
        joined = " ".join(cell.strip() for cell in row if cell.strip())
        if _meaningful(joined):
            items.append(joined)
    return text, items


def _from_docx(data: bytes) -> tuple[str, list[str]]:
    from docx import Document

    _guard_zip_bomb(data)
    document = Document(io.BytesIO(data))
    items = [p.text.strip() for p in document.paragraphs if _meaningful(p.text)]
    for table in document.tables:
        for row in table.rows:
            joined = " ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if _meaningful(joined):
                items.append(joined)
    return "\n".join(items), items


def _from_xlsx(data: bytes) -> tuple[str, list[str]]:
    from openpyxl import load_workbook

    _guard_zip_bomb(data)
    workbook = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    items: list[str] = []
    for sheet in workbook.worksheets:
        for row in sheet.iter_rows(values_only=True):
            cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
            joined = " ".join(cells)
            if _meaningful(joined):
                items.append(joined)
    workbook.close()
    return "\n".join(items), items


def _from_pdf(data: bytes) -> tuple[str, list[str]]:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    if len(reader.pages) > _MAX_PDF_PAGES:
        raise DocumentTooLarge(f"PDF çok fazla sayfa içeriyor ({len(reader.pages)})")
    chunks: list[str] = []
    total = 0
    for page in reader.pages:
        chunk = page.extract_text() or ""
        chunks.append(chunk)
        total += len(chunk)
        if total > _MAX_PDF_TEXT_CHARS:
            raise DocumentTooLarge("PDF çıkarılan metni güvenlik sınırını aştı")
    return _from_text("\n".join(chunks))


def parse_document(filename: str, data: bytes) -> tuple[str, list[str]]:
    ext = Path(filename).suffix.lower()
    if ext == ".csv":
        return _from_csv(data, ",")
    if ext == ".tsv":
        return _from_csv(data, "\t")
    if ext == ".docx":
        return _from_docx(data)
    if ext == ".xlsx":
        return _from_xlsx(data)
    if ext == ".pdf":
        return _from_pdf(data)
    # .txt / .md / unknown → plain text
    return _from_text(data.decode("utf-8", errors="replace"))
